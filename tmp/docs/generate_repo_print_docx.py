from __future__ import annotations

import os
import re
from datetime import datetime
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt
from PIL import Image


ROOT = Path(__file__).resolve().parents[2]
OUTPUT = ROOT / "output" / "doc" / "repo_printable.docx"
TMP_IMAGES = ROOT / "tmp" / "docs" / "normalized_images"

TEXT_SUFFIXES = {".md", ".txt", ".py", ".cpp", ".h", ".ini", ".yaml", ".yml", ".json"}
IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg"}
TEXT_EXPLICIT = {
    ".gitignore",
    "LICENSE",
    "src/include/README",
    "src/lib/README",
    "src/test/README",
}
SKIP_DIRS = {".git", "__pycache__", ".venv", "venv", "node_modules", "output", "tmp"}


def iter_files() -> tuple[list[Path], list[Path], list[Path]]:
    text_files: list[Path] = []
    image_files: list[Path] = []
    other_files: list[Path] = []

    for path in sorted(ROOT.rglob("*")):
        if not path.is_file():
            continue
        rel = path.relative_to(ROOT)
        if any(part in SKIP_DIRS for part in rel.parts):
            continue
        rel_str = rel.as_posix()
        if path.suffix.lower() in TEXT_SUFFIXES or rel_str in TEXT_EXPLICIT:
            text_files.append(path)
        elif path.suffix.lower() in IMAGE_SUFFIXES:
            image_files.append(path)
        else:
            other_files.append(path)
    return text_files, image_files, other_files


def ensure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(11)

    if "Code Block" not in doc.styles:
        style = doc.styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = doc.styles["Normal"]
        style.font.name = "Courier New"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
        style.font.size = Pt(9)


def set_page_layout(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    add_footer(section)


def add_footer(section) -> None:
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.text = "KU STEAM Pinkies repository printout | Page "
    add_page_field(footer_para)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_begin, instr, fld_char_end])


def add_title_page(doc: Document, text_count: int, image_count: int, other_count: int) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("KU STEAM Pinkies\nRepository Printout")
    run.bold = True
    run.font.size = Pt(20)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Prepared for printed review in DOCX format").italic = True

    meta = [
        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}",
        f"Repository root: {ROOT.name}",
        f"Text files included: {text_count}",
        f"Images embedded: {image_count}",
        f"Other binary/reference files listed: {other_count}",
    ]
    for line in meta:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.add_run(line)

    doc.add_paragraph("")
    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    intro.add_run(
        "This document collects the repository tree, text documentation, source files, and printable image evidence "
        "from the WRO 2026 Future Engineers submission repository."
    )
    doc.add_page_break()


def add_repo_tree(doc: Document, all_files: Iterable[Path]) -> None:
    doc.add_heading("Repository Tree", level=1)
    for path in all_files:
        rel = path.relative_to(ROOT).as_posix()
        para = doc.add_paragraph(style="Code Block")
        para.paragraph_format.space_after = Pt(0)
        para.add_run(rel)
    doc.add_page_break()


def add_markdown_or_text(doc: Document, path: Path) -> None:
    rel = path.relative_to(ROOT).as_posix()
    doc.add_heading(rel, level=1)

    content = path.read_text(encoding="utf-8", errors="replace").splitlines()
    suffix = path.suffix.lower()
    code_mode = suffix in {".py", ".cpp", ".h", ".ini", ".yaml", ".yml", ".json"} or rel.startswith("src/")

    if code_mode:
        for line in content:
            para = doc.add_paragraph(style="Code Block")
            para.paragraph_format.space_after = Pt(0)
            para.add_run(line.expandtabs(4) if line else " ")
        doc.add_page_break()
        return

    in_fence = False
    for raw_line in content:
        line = raw_line.rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            in_fence = not in_fence
            continue

        if in_fence:
            para = doc.add_paragraph(style="Code Block")
            para.paragraph_format.space_after = Pt(0)
            para.add_run(line.expandtabs(4) if line else " ")
            continue

        if not stripped:
            doc.add_paragraph("")
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=4)
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=3)
            continue
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=2)
            continue
        if stripped.startswith("- "):
            para = doc.add_paragraph(style="List Bullet")
            para.add_run(stripped[2:].strip())
            continue
        if re.match(r"^\d+\.\s+", stripped):
            para = doc.add_paragraph(style="List Number")
            para.add_run(re.sub(r"^\d+\.\s+", "", stripped))
            continue

        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run(line)

    doc.add_page_break()


def normalized_image_path(image: Path) -> Path:
    TMP_IMAGES.mkdir(parents=True, exist_ok=True)
    target = TMP_IMAGES / (image.stem.replace(" ", "_") + ".png")
    with Image.open(image) as img:
        safe = img.convert("RGB")
        safe.save(target, format="PNG")
    return target


def add_images(doc: Document, image_files: list[Path]) -> None:
    if not image_files:
        return
    doc.add_heading("Image Appendix", level=1)
    for image in image_files:
        rel = image.relative_to(ROOT).as_posix()
        doc.add_heading(rel, level=2)
        try:
            doc.add_picture(str(image), width=Inches(6.2))
        except Exception:
            doc.add_picture(str(normalized_image_path(image)), width=Inches(6.2))
        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = caption.add_run(rel)
        cap_run.italic = True
    doc.add_page_break()


def add_other_files(doc: Document, other_files: list[Path]) -> None:
    if not other_files:
        return
    doc.add_heading("Other Repository Files", level=1)
    intro = doc.add_paragraph()
    intro.add_run(
        "These files were not embedded as full printable text. They remain referenced here as part of the repository artifact set."
    )
    for path in other_files:
        rel = path.relative_to(ROOT).as_posix()
        para = doc.add_paragraph(style="List Bullet")
        para.add_run(rel)


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)

    text_files, image_files, other_files = iter_files()
    all_files = sorted(text_files + image_files + other_files)

    doc = Document()
    ensure_styles(doc)
    set_page_layout(doc)
    add_title_page(doc, len(text_files), len(image_files), len(other_files))
    add_repo_tree(doc, all_files)

    doc.add_heading("Text and Source Files", level=1)
    for path in text_files:
        add_markdown_or_text(doc, path)

    add_images(doc, image_files)
    add_other_files(doc, other_files)
    doc.save(str(OUTPUT))
    print(OUTPUT)


if __name__ == "__main__":
    main()
