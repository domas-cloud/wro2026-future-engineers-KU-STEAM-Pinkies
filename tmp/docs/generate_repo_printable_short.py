from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image


ROOT = Path(__file__).resolve().parents[2]
OUTPUT = ROOT / "output" / "doc" / "judge_packet_standalone.docx"
TMP_IMAGES = ROOT / "tmp" / "docs" / "normalized_images_short"

def set_font(run, size: int, bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)


def style_document(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(11)

    for name, size in [("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11)]:
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.9)
    section.right_margin = Cm(1.9)
    add_footer(section)


def add_footer(section) -> None:
    para = section.footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run("KU STEAM Pinkies | Judge Report | Page ")
    set_font(run, 9)
    add_page_field(para)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_begin, instr, fld_char_end])


def add_rule(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "808080")
    borders.append(bottom)


def safe_add_picture(run, image: Path, width_inches: float) -> None:
    try:
        run.add_picture(str(image), width=Inches(width_inches))
        return
    except Exception:
        TMP_IMAGES.mkdir(parents=True, exist_ok=True)
        normalized = TMP_IMAGES / f"{image.stem.replace(' ', '_')}.png"
        with Image.open(image) as img:
            img.convert("RGB").save(normalized, format="PNG")
        run.add_picture(str(normalized), width=Inches(width_inches))


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("KU STEAM Pinkies")
    set_font(r, 20, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("WRO 2026 Future Engineers")
    set_font(r, 16, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Judge Report")
    set_font(r, 13, italic=True)

    doc.add_paragraph("")

    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = intro.add_run(
        "This document is intended to work on its own. "
        "It summarizes the robot design, hardware, software, testing, and engineering decisions without requiring any extra reference material."
    )
    set_font(r, 11)

    meta = doc.add_table(rows=4, cols=2)
    meta.style = "Table Grid"
    rows = [
        ("Team", "KU STEAM Pinkies"),
        ("Competition", "WRO 2026 Future Engineers"),
        ("Robot concept", "Compact autonomous car with split perception and control"),
        ("Document purpose", "Self-contained printed engineering summary"),
    ]
    for i, (k, v) in enumerate(rows):
        meta.cell(i, 0).text = k
        meta.cell(i, 1).text = v

    doc.add_paragraph("")
    img = ROOT / "v-photos" / "front.jpg"
    if img.exists():
        pic = doc.add_paragraph()
        pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
        safe_add_picture(pic.add_run(), img, 5.8)
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run("Front view of the final robot")
        set_font(r, 10, italic=True)

    doc.add_page_break()


def add_section_title(doc: Document, title: str, subtitle: str | None = None) -> None:
    p = doc.add_paragraph()
    p.style = "Heading 1"
    r = p.add_run(title)
    set_font(r, 16, bold=True)
    add_rule(p)
    if subtitle:
        s = doc.add_paragraph()
        s.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = s.add_run(subtitle)
        set_font(r, 11)


def add_subsection_title(doc: Document, title: str) -> None:
    p = doc.add_paragraph()
    p.style = "Heading 2"
    r = p.add_run(title)
    set_font(r, 13, bold=True)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(item)
        set_font(r, 11)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(item)
        set_font(r, 11)


def add_paragraph(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_prefix:
        r = p.add_run(bold_prefix)
        set_font(r, 11, bold=True)
    r = p.add_run(text)
    set_font(r, 11)


def add_photo_row(doc: Document, photos: list[tuple[Path, str]], width_inches: float = 2.8) -> None:
    table = doc.add_table(rows=1, cols=len(photos))
    table.style = "Table Grid"
    for i, (img, caption) in enumerate(photos):
        cell = table.cell(0, i)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if img.exists():
            safe_add_picture(p.add_run(), img, width_inches)
        cp = cell.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cp.add_run(caption)
        set_font(r, 9, italic=True)


def add_robot_gallery(doc: Document) -> None:
    add_section_title(
        doc,
        "1. Robot Gallery",
        "This section shows the existing final robot from several angles so the printed version includes direct visual evidence, not only text descriptions.",
    )
    add_photo_row(
        doc,
        [
            (ROOT / "v-photos" / "front.jpg", "Front view"),
            (ROOT / "v-photos" / "left.jpg", "Left side"),
        ],
    )
    add_photo_row(
        doc,
        [
            (ROOT / "v-photos" / "right.jpg", "Right side"),
            (ROOT / "v-photos" / "back.jpg", "Rear view"),
        ],
    )
    add_photo_row(
        doc,
        [
            (ROOT / "v-photos" / "top.jpg", "Top view"),
            (ROOT / "v-photos" / "bottom.jpg", "Bottom view"),
        ],
    )
    add_photo_row(
        doc,
        [
            (ROOT / "t-photos" / "oficial.jpg", "Team and robot photo"),
            (ROOT / "t-photos" / "funny.jpg", "Additional team photo"),
        ],
        width_inches=2.6,
    )
    doc.add_page_break()


def add_key_specs(doc: Document) -> None:
    add_section_title(
        doc,
        "2. Robot Summary",
        "The final robot was designed as a compact autonomous driving system where mechanics, sensing, and control support each other rather than being optimized separately.",
    )
    add_paragraph(
        doc,
        "The robot was documented to show engineering decisions clearly. Instead of presenting one isolated subsystem as the main success, the final build focuses on how chassis geometry, steering, drivetrain, sensing, and control logic interact during repeated runs on the field.",
    )
    add_bullets(
        doc,
        [
            "Rear-wheel drive and front-wheel steering layout.",
            "ESP32 for real-time low-level control and a Raspberry Pi Zero with camera for perception.",
            "BNO085 IMU plus three VL53L4CD distance sensors for local geometry and heading control.",
            "MG90S steering servo, N20 6 V 600 rpm drive motor, and L298N motor driver.",
            "Small chassis chosen for repeatability, turning efficiency, and easier packaging.",
        ],
    )
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Subsystem"
    table.rows[0].cells[1].text = "Final choice"
    data = [
        ("Drive", "Rear-wheel drive with LEGO differential"),
        ("Steering", "Front steering, refined V2/V3 geometry"),
        ("Compute", "Raspberry Pi Zero + ESP32 split architecture"),
        ("Sensors", "Camera, BNO085, 3x VL53L4CD"),
        ("Power", "2x 18650 pack with separated regulated branches"),
    ]
    for left, right in data:
        row = table.add_row().cells
        row[0].text = left
        row[1].text = right

    add_subsection_title(doc, "Main Design Goal")
    add_paragraph(
        doc,
        "The main goal was not to maximize any single specification. The goal was to build a compact robot that remains predictable, tunable, and repeatable in straight driving, corner handling, and obstacle-related transitions.",
    )
    add_subsection_title(doc, "Physical Dimensions")
    add_bullets(
        doc,
        [
            "Approximate length: 21 cm.",
            "Approximate width: 10 cm.",
            "Approximate height: 8 cm.",
        ],
    )


def add_architecture(doc: Document) -> None:
    add_section_title(doc, "3. System Architecture")
    add_paragraph(
        doc,
        "The architecture is intentionally split into two layers. The Raspberry Pi Zero and camera interpret the wider scene and choose a driving reference, while the ESP32 handles the time-critical loop: sensor polling, steering output, motor output, and turn execution.",
    )
    add_bullets(
        doc,
        [
            "Perception layer: selects path or obstacle side.",
            "Low-level control layer: reads yaw and distance sensors, applies steering correction, and executes turns.",
            "This split kept the controller simple, fast, and easier to tune on the real field.",
        ],
    )
    add_subsection_title(doc, "Low-Level Runtime Flow")
    add_numbered(
        doc,
        [
            "Read button, yaw, and local distance sensors.",
            "Wait safely before the run starts.",
            "Keep heading and side distance under control on straight sections.",
            "Trigger the hard-turn routine when the front threshold is reached.",
            "Count sectors and stop at the required final state.",
        ],
    )

    img = ROOT / "schemes" / "images" / "schematic-overview.png"
    if img.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        safe_add_picture(p.add_run(), img, 5.9)
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run("Simplified electronics and wiring overview")
        set_font(r, 10, italic=True)


def add_mechanics(doc: Document) -> None:
    add_section_title(doc, "4. Mechanical Design")
    add_paragraph(
        doc,
        "Mechanical development concentrated on stability and motion quality rather than on visual complexity. The smaller final chassis improved packaging, turning behaviour, and repeatability, while the refined steering geometry reduced wasted force and improved how reliably the front axle followed the commanded angle.",
    )
    add_bullets(
        doc,
        [
            "Compact chassis for easier manoeuvrability and more practical packaging.",
            "Rear axle with differential to reduce binding in corners.",
            "Front steering geometry improved across versions to lower servo load.",
            "Silicone front wheels selected because grip mattered more than matching wheel type on both axles.",
        ],
    )

    img = ROOT / "docs" / "design" / "images" / "steering-v3.png"
    if img.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        safe_add_picture(p.add_run(), img, 4.8)
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run("Final steering concept")
        set_font(r, 10, italic=True)


def add_electronics(doc: Document) -> None:
    add_section_title(doc, "5. Electronics And Sensors")
    add_paragraph(
        doc,
        "The electronics layout was split on purpose. The camera side handled perception and the ESP32 side handled control. Power distribution was also separated into practical branches so motor and servo disturbances would not easily affect logic and sensing.",
    )
    add_bullets(
        doc,
        [
            "Raspberry Pi Zero for perception and camera-side decisions.",
            "ESP32-WROOM-32 for sensor polling, steering, and motor output.",
            "BNO085 for heading reference and orientation stability.",
            "Three VL53L4CD sensors for front and side distance feedback.",
            "2x 18650 battery pack with separated regulated branches.",
        ],
    )
    add_subsection_title(doc, "Why This Layout Worked")
    add_bullets(
        doc,
        [
            "The low-level controller stayed responsive and easier to debug.",
            "Sensor and logic rails had better protection from motor noise and voltage sag.",
            "The full system remained easier to explain and rebuild from documentation.",
        ],
    )
    add_subsection_title(doc, "Power And Wiring Structure")
    add_bullets(
        doc,
        [
            "2x 18650 battery pack feeds the motor branch, logic branches, sensor rail, and servo branch.",
            "The L298N and N20 motor stay on the higher-current drive path.",
            "The ESP32, Raspberry Pi Zero, IMU, and ToF sensors stay on regulated logic or sensor rails.",
            "The servo branch is kept separate so steering current spikes do not disturb the rest of the system.",
            "All subsystems share a common ground reference.",
        ],
    )


def add_software(doc: Document) -> None:
    add_section_title(doc, "6. Software Logic")
    add_paragraph(
        doc,
        "The low-level software evidence is the ESP32 control loop itself. It shows the practical runtime behavior used on the robot: initialization, waiting for the run start, reading yaw and distance sensors, holding heading and wall offset, executing hard turns, and finishing after the required sequence.",
    )
    add_bullets(
        doc,
        [
            "Startup initializes the I2C devices, IMU, distance sensors, PWM, and motor output.",
            "The controller waits for the start button before motion begins.",
            "Straight sections use heading and side-distance correction together.",
            "Corners are handled by a dedicated hard-turn routine when the front threshold is reached.",
            "The perception layer is intended to guide the reference line rather than replace the low-level loop.",
        ],
    )
    add_subsection_title(doc, "Low-Level Control Law")
    add_paragraph(
        doc,
        "On straight sections, the steering command combines heading error, side-distance error, and a damping term. In simple form, the robot computes a correction from yaw and wall offset, clamps the final result into the allowed servo range, and keeps driving until the front sensor detects a corner trigger.",
    )
    add_subsection_title(doc, "Obstacle Rule")
    add_bullets(
        doc,
        [
            "Red pillar means pass on the right side.",
            "Green pillar means pass on the left side.",
            "The clean implementation principle is to shift the reference line while keeping the same low-level steering controller.",
        ],
    )


def add_engineering(doc: Document) -> None:
    add_section_title(doc, "7. Main Engineering Decisions")
    add_paragraph(
        doc,
        "The strongest engineering pattern in the project was choosing the version that improved repeatability, not the version that looked stronger on paper. Several final decisions therefore favored precision, recoverability, and ease of tuning instead of extreme parameters.",
    )
    add_bullets(
        doc,
        [
            "A smaller robot was chosen instead of a larger and mechanically heavier concept.",
            "The 600 rpm N20 motor gave a better speed-torque balance than 300 rpm or 1000 rpm options.",
            "Steering range was limited to about 60 degrees because maximum angle was less stable in practice.",
            "Steering geometry was corrected instead of replacing the servo with a stronger one.",
            "Silicone front wheels improved real steering effect by reducing slip.",
            "The LEGO differential produced smoother and more repeatable cornering than the earlier metal version.",
        ],
    )

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Problem"
    table.rows[0].cells[1].text = "Decision"
    table.rows[0].cells[2].text = "Reason"
    rows = [
        ("Servo load", "Fix steering geometry", "Removed bad lever arm and improved repeatability"),
        ("Wheel slip", "Use silicone front wheels", "Improved conversion of steering command into motion"),
        ("Corner binding", "Use LEGO differential", "Smoother turns and less resistance"),
        ("Motor extremes", "Use 600 rpm N20", "Balanced speed and useful torque"),
    ]
    for a, b, c in rows:
        cells = table.add_row().cells
        cells[0].text = a
        cells[1].text = b
        cells[2].text = c


def add_testing(doc: Document) -> None:
    add_section_title(doc, "8. Testing And Performance")
    add_paragraph(
        doc,
        "Testing was used to compare versions under repeated practical runs instead of relying on isolated bench impressions. The final design was selected because it was easier to control, more repeatable, and less sensitive to repeated failure patterns.",
    )
    add_bullets(
        doc,
        [
            "Straight-drive drift over about 3 m.",
            "Space required for a 90-degree turn.",
            "Steering return-to-center repeatability.",
            "Left-right symmetry and stability across repeated runs.",
            "Track behaviour near obstacles, reflective surfaces, and after corrections.",
        ],
    )
    add_subsection_title(doc, "Observed Conclusions")
    add_bullets(
        doc,
        [
            "Corrected steering geometry reduced servo load and improved center stability.",
            "The LEGO differential reduced turning resistance and improved corner smoothness.",
            "The 600 rpm motor gave the best balance between speed and usable torque.",
            "More rigid sensor and IMU mounting improved consistency across runs.",
        ],
    )

    img1 = ROOT / "docs" / "design" / "images" / "metal-differential.jpg"
    img2 = ROOT / "docs" / "design" / "images" / "lego-differential.png"
    if img1.exists() and img2.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        safe_add_picture(p.add_run(), img1, 2.7)
        p.add_run("  ")
        safe_add_picture(p.add_run(), img2, 2.7)
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run("Earlier metal differential and final LEGO differential")
        set_font(r, 10, italic=True)


def add_evaluation(doc: Document) -> None:
    add_section_title(doc, "9. Evaluation, Lessons, And Limits")
    add_subsection_title(doc, "What Worked")
    add_bullets(
        doc,
        [
            "The corrected steering geometry reduced servo load and improved repeatability.",
            "The differential remained a necessary drivetrain element for smoother corners.",
            "A mixed sensor strategy worked better than relying on only one sensor type.",
            "Rigid IMU mounting improved heading reliability during repeated runs.",
            "The simpler split architecture made the full robot easier to tune and explain.",
        ],
    )
    add_subsection_title(doc, "What Did Not Work")
    add_bullets(
        doc,
        [
            "Early steering with a large wheel lever arm overloaded the servo and reduced reliability.",
            "Driving without a suitable differential increased turning resistance and slip.",
            "Over-reliance on a single sensor type was not stable enough across scenarios.",
            "Weak or flexible mounting reduced sensor trustworthiness.",
        ],
    )
    add_subsection_title(doc, "Current Limits")
    add_bullets(
        doc,
        [
            "The current documentation does not yet present a fully laboratory-style numeric dataset.",
            "Final performance is documented mainly through repeated practical comparison and observed field behavior.",
            "The strongest future addition would be a stricter counted dataset for drift, obstacle runs, and success rate.",
        ],
    )
    add_subsection_title(doc, "What A Judge Should Take From This")
    add_paragraph(
        doc,
        "The final robot was not selected because it had the strongest individual part. It was selected because the whole system became more repeatable after the steering geometry, differential behavior, front grip, sensor stability, and controller tuning started working together.",
    )


def add_risks_and_safety(doc: Document) -> None:
    add_section_title(doc, "10. Risks And Safety")
    add_paragraph(
        doc,
        "Risk analysis was treated as an engineering tool, not as an afterthought. The main idea was to identify the situations that most often reduced repeatability and then redesign the robot to remove those root causes instead of only adding stronger parts.",
    )
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Risk"
    table.rows[0].cells[1].text = "Why it mattered"
    table.rows[0].cells[2].text = "Mitigation"
    rows = [
        ("Large steering lever arm", "High servo load and weak steering efficiency", "Redesigned steering geometry"),
        ("Too much steering angle", "Lower stability", "Limited usable steering range"),
        ("Front-wheel slip", "Weak real steering effect", "Changed wheel material and geometry"),
        ("Poor differential behavior", "Binding and lower precision", "Kept a better differential solution"),
        ("Wrong motor balance", "Too slow or too weak under load", "Compared 300/600/1000 rpm options"),
        ("Straight-driving drift", "Lower repeatability", "Improved multiple connected mechanical factors"),
    ]
    for a, b, c in rows:
        cells = table.add_row().cells
        cells[0].text = a
        cells[1].text = b
        cells[2].text = c

    add_subsection_title(doc, "Visible Safety Logic")
    add_bullets(
        doc,
        [
            "The robot does not start moving before the start button is pressed.",
            "If a critical sensor fails at startup, the controller halts instead of starting blindly.",
            "Steering output is clamped to the safe mechanical range.",
            "Distance-based corrections are used only inside practical trust windows.",
            "The robot stops cleanly after the run logic is complete.",
        ],
    )


def add_reproducibility(doc: Document) -> None:
    add_section_title(doc, "11. Reproducibility And Reading Path")
    add_paragraph(
        doc,
        "This printed report is meant to be usable without any extra files. For that reason, the most important rebuild information is summarized directly here instead of assuming access to outside material.",
    )
    add_subsection_title(doc, "What Another Team Would Need To Rebuild")
    add_numbered(
        doc,
        [
            "A compact chassis with rear-wheel drive and front-wheel steering.",
            "A rear differential and the final low-resistance steering geometry.",
            "An ESP32 low-level controller and a Raspberry Pi Zero with camera for perception.",
            "A BNO085 IMU plus three VL53L4CD sensors mounted as front, left, and right.",
            "A 2x 18650 power system with separated motor, servo, logic, and sensor branches.",
            "The control logic described in this document: heading hold, side-distance correction, and hard-turn corner handling.",
        ],
    )


def add_submission_assets(doc: Document) -> None:
    add_section_title(doc, "12. Submission Assets And Build Evidence")
    add_bullets(
        doc,
        [
            "Robot photos from the front, sides, rear, top, and bottom are included in this report.",
            "The document explains the electrical split between the Raspberry Pi Zero, ESP32, sensors, servo, motor driver, and battery pack.",
            "The key control logic is described directly in words, so the robot behavior can be understood even without source code access.",
            "The main design changes, test results, and rejected ideas are all summarized here.",
            "This makes the report usable as a stand-alone engineering explanation of the final robot.",
        ],
    )
    add_paragraph(
        doc,
        "A good WRO explanation should show both the physical robot and the reasoning behind it. That is why this report includes photographs, subsystem descriptions, test-based decisions, and practical rebuild information in one place.",
    )


def add_component_list(doc: Document) -> None:
    add_section_title(doc, "13. Main Component List")
    add_paragraph(
        doc,
        "The table below summarizes the most important hardware used in the final robot so the reader can understand the build without referring to an external parts list.",
    )
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Subsystem"
    table.rows[0].cells[1].text = "Main components"
    rows = [
        ("Compute and control", "Raspberry Pi Zero, ESP32-WROOM-32"),
        ("Sensing", "Camera module, BNO085 IMU, 3x VL53L4CD"),
        ("Drive", "N20 6 V 600 rpm motor, L298N motor driver"),
        ("Steering", "MG90S servo, three-gear steering mechanism"),
        ("Power", "2x 18650 Li-ion pack, step-down regulation, separated rails"),
        ("Mechanics", "Wooden frame, custom mounts, 3D-printed steering-related parts, rear differential"),
    ]
    for left, right in rows:
        cells = table.add_row().cells
        cells[0].text = left
        cells[1].text = right


def add_standalone_summary(doc: Document) -> None:
    add_section_title(doc, "14. Standalone Summary For Judges")
    add_paragraph(
        doc,
        "If the reader sees only this report, the key conclusion is still clear: the robot is a compact autonomous car built around repeatability. Its final performance comes from a balanced combination of chassis size, steering geometry, differential behavior, sensor roles, and a simple two-layer control architecture.",
    )
    add_bullets(
        doc,
        [
            "Mechanically, the decisive improvements were steering geometry, front-wheel grip, wheel mounting quality, and the LEGO differential.",
            "Electronically, the decisive improvement was keeping perception and low-level control separate while protecting logic and sensors from noisy power behavior.",
            "In software, the decisive choice was to keep the low-level controller simple: heading hold, side-distance correction, and discrete hard turns.",
            "In testing, the decisive rule was to keep the version that repeated well, not the version that only looked stronger in theory.",
        ],
    )


def add_closing(doc: Document) -> None:
    add_section_title(doc, "15. Closing Note")
    add_paragraph(
        doc,
        "This report is intentionally written as a stand-alone explanation of the robot. A judge should be able to understand what was built, why the main decisions were made, how the control system works, what was tested, and how the robot could be rebuilt, even without opening any additional material.",
    )


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_document(doc)
    add_cover(doc)
    add_robot_gallery(doc)
    add_key_specs(doc)
    add_architecture(doc)
    add_mechanics(doc)
    add_electronics(doc)
    add_software(doc)
    add_engineering(doc)
    add_testing(doc)
    add_evaluation(doc)
    add_risks_and_safety(doc)
    add_reproducibility(doc)
    add_submission_assets(doc)
    add_component_list(doc)
    add_standalone_summary(doc)
    add_closing(doc)
    doc.save(str(OUTPUT))
    print(OUTPUT)


if __name__ == "__main__":
    main()
